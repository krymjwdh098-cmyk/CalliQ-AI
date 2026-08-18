"""
CV Parser — text extraction + AI parsing (provider-agnostic).
Fix: every parse is strictly isolated, no state shared between calls.
"""
import hashlib
import logging
import time
from pathlib import Path
from typing import Optional

from service.llm_router import llm_complete_json, llm_complete_with_fallback
import json as _json
import re


def _parse_json(raw: str, agent: Optional[str] = None, user: str = "") -> dict:
    """Parse JSON from LLM response"""
    text = raw.strip()
    # Strip markdown fences
    if "```" in text:
        parts = text.split("```")
        for part in parts:
            part = part.strip()
            if part.startswith("json"):
                part = part[4:].strip()
            try:
                return _json.loads(part)
            except Exception:
                continue
    # Direct parse
    try:
        return _json.loads(text)
    except _json.JSONDecodeError:
        # Try to find JSON object in text
        match = re.search(r'\{.*\}', text, re.DOTALL)
        if match:
            try:
                return _json.loads(match.group())
            except Exception:
                pass
    logger.error(f"JSON parse failed. Raw preview: {raw[:300]}")
    return {}

logger = logging.getLogger(__name__)


# ── Prompts ───────────────────────────────────────────────────────────────────

PARSE_SYSTEM = """You are an expert HR data extraction AI.
Extract ALL information from the CV. Return ONLY valid JSON — no markdown, no code fences.

EXTRACTION RULES:
- years_experience: calculate from work history dates, NOT self-reported. If dates missing, estimate.
- salary_expectation: extract ONLY if explicitly in CV (e.g. "Expected: $5000/month"). Else null.
- notice_period_days: extract if stated ("1 month notice"→30, "2 weeks"→14, "immediate"→0). Else null.
- availability_date: "immediately"→"immediate", "June 2025"→"2025-06-01", else null.
- For certifications: extract name, issuer, year if available.
- For previous_positions: estimate duration_months from dates if possible.
- Return null for missing fields, never empty string for nullable fields.

Return this exact JSON:
{
  "full_name": "string",
  "email": "string or null",
  "phone": "string or null",
  "location": "string or null",
  "nationality": "string or null",
  "linkedin": "string or null",
  "github": "string or null",
  "portfolio": "string or null",
  "current_position": "string or null",
  "current_company": "string or null",
  "years_experience": 0.0,
  "previous_positions": [{"title":"","company":"","start":"","end":"","duration_months":0}],
  "companies": ["string"],
  "education": [{"degree":"","field":"","institution":"","year":"","gpa":null}],
  "certifications": [{"name":"","issuer":"","year":null}],
  "courses": ["string"],
  "technical_skills": {"category": ["skill"]},
  "soft_skills": ["string"],
  "languages": [{"language":"","level":"Native|Fluent|Professional|Conversational|Basic"}],
  "projects": [{"name":"","description":"","technologies":[],"url":null}],
  "achievements": ["string with measurable impact if available"],
  "awards": ["string"],
  "salary_expectation": null,
  "salary_currency": null,
  "notice_period_days": null,
  "availability_date": null,
  "remote_preference": null,
  "ai_summary": "2-3 sentence professional summary"
}"""

MATCH_SYSTEM = """You are a senior AI recruiter scoring a candidate against a specific job.
Return ONLY valid JSON — no markdown, no code fences, no explanation.

SCORING WEIGHTS (must be applied exactly):
  skill_match        → 35% of overall_score   (required skills coverage)
  experience_match   → 25% of overall_score   (years + relevance of experience)
  education_match    → 15% of overall_score   (degree, field, institution level)
  seniority_match    → 10% of overall_score   (current title vs job level)
  keyword_match      → 10% of overall_score   (ATS keyword density from job description)
  location_match     →  5% of overall_score   (remote-ok or location fit)

overall_score = (skill_match*0.35 + experience_match*0.25 + education_match*0.15
                 + seniority_match*0.10 + keyword_match*0.10 + location_match*0.05)

SALARY MATCH (separate, does not affect overall_score):
  salary_match: 0-100 based on candidate expectation vs job budget.
  If either is unknown, return 50 (neutral).

RECOMMENDATION THRESHOLDS:
  overall_score >= 80  → "Strong Hire"
  overall_score >= 65  → "Hire"
  overall_score >= 50  → "Consider"
  overall_score < 50   → "Reject"

ATS SCORE (separate technical assessment):
  ats_score: 0-100 based on keyword density, formatting signals, and completeness.
  ats_issues: specific problems found (e.g. "Missing LinkedIn", "No quantified achievements").
  ats_suggestions: actionable fixes for the candidate.

Return this exact JSON schema:
{
  "overall_score": 0.0,
  "skill_match": 0.0,
  "experience_match": 0.0,
  "education_match": 0.0,
  "seniority_match": 0.0,
  "location_match": 0.0,
  "keyword_match": 0.0,
  "salary_match": 0.0,
  "ai_confidence": 0.0,
  "recommendation": "Strong Hire|Hire|Consider|Reject",
  "recommendation_reason": "2-3 sentences with specific evidence from the CV",
  "ai_summary": "2-3 sentence executive summary for the recruiter",
  "strengths": ["specific strength with evidence"],
  "weaknesses": ["specific gap with impact"],
  "missing_skills": ["skill name"],
  "missing_certs": ["certification name"],
  "matched_skills": ["skill name"],
  "matched_requirements": ["requirement text"],
  "missing_requirements": ["requirement text"],
  "skill_gap_analysis": "paragraph explaining exact gaps and how critical they are",
  "ats_score": 0.0,
  "ats_issues": ["specific issue"],
  "ats_suggestions": ["specific actionable fix"],
  "notice_period_days": null,
  "salary_expectation_match": "within_range|above_range|below_range|unknown"
}

Be objective. Base every score on ACTUAL CV content vs ACTUAL job requirements.
Do not inflate scores. A candidate missing 3/5 required skills must score below 60."""


# ── Text Extraction ───────────────────────────────────────────────────────────

def compute_file_hash(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


def extract_text_from_file(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()
    logger.info(f"Extracting text from {file_path} (type={suffix})")
    t0 = time.time()

    if suffix == ".pdf":
        text = _extract_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        text = _extract_docx(file_path)
    elif suffix in (".jpg", ".jpeg", ".png"):
        text = _extract_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    elapsed = int((time.time() - t0) * 1000)
    logger.info(f"Text extracted: {len(text)} chars in {elapsed}ms")
    return text


def _extract_pdf(path: str) -> str:
    # Strategy 1: pypdf (fast, works for text PDFs)
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages = []
        for page in reader.pages:
            t = page.extract_text() or ""
            if t.strip():
                pages.append(t)
        text = "\n".join(pages).strip()
        if len(text) > 100:  # meaningful content
            return text
        logger.info("pypdf extracted too little text, trying pdfplumber")
    except Exception as e:
        logger.warning(f"pypdf failed: {e}")

    # Strategy 2: pdfplumber (better layout handling)
    try:
        import importlib
        pdfplumber = importlib.import_module("pdfplumber")
        with pdfplumber.open(path) as pdf:
            pages = []
            for page in pdf.pages:
                t = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                if t.strip():
                    pages.append(t)
        text = "\n".join(pages).strip()
        if len(text) > 100:
            return text
        logger.info("pdfplumber extracted too little, trying OCR")
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")

    # Strategy 3: OCR for scanned PDFs
    return _ocr_pdf(path)


def _ocr_pdf(path: str) -> str:
    """OCR fallback for scanned/image PDFs."""
    try:
        # Try pdf2image + tesseract
        import importlib
        pdf2image = importlib.import_module("pdf2image")
        images = pdf2image.convert_from_path(path, dpi=200, first_page=1, last_page=3)
        texts = []
        for img in images:
            t = _extract_image_ocr_pil(img)
            if t.strip():
                texts.append(t)
        text = "\n".join(texts).strip()
        if text:
            logger.info(f"OCR extracted {len(text)} chars from scanned PDF")
            return text
    except ImportError:
        logger.warning("pdf2image not installed — install with: pip install pdf2image")
    except Exception as e:
        logger.warning(f"pdf2image OCR failed: {e}")

    # Last resort: extract any embedded text + signal it's likely scanned
    return "[Scanned PDF — text extraction limited. Manual review recommended.]"


def _extract_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def _extract_image_ocr_pil(img) -> str:
    """OCR from PIL Image object (in-memory, used by pdf OCR too)."""
    try:
        import pytesseract
        if hasattr(img, "mode") and img.mode != "L":
            img = img.convert("L")
        return pytesseract.image_to_string(img, lang="eng+ara")
    except ImportError:
        raise ImportError("Install pytesseract: pip install pytesseract")


def _extract_image(path: str) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(path)
        if img.mode != "L":
            img = img.convert("L")
        return pytesseract.image_to_string(img, lang="eng+ara")
    except ImportError:
        raise ImportError("Install pytesseract and Pillow for image OCR")


# ── AI Parsing ────────────────────────────────────────────────────────────────

def parse_cv_with_ai(raw_text: str) -> dict:
    """
    Parse CV text with AI. Returns structured dict.
    Uses Groq first, fallback to Gemini, returns empty if both fail.
    """
    if not raw_text or len(raw_text.strip()) < 20:
        logger.warning("CV text too short for AI parsing")
        return {"full_name": "Unknown", "email": None, "phone": None}

    t0 = time.time()
    
    # Use fallback system: Groq → Gemini → empty
    raw_result = llm_complete_with_fallback(
        system=PARSE_SYSTEM,
        user=f"CV Text:\n\n{raw_text[:12000]}",
        max_tokens=2000,
        agent="parser",
    )
    
    result = _parse_json(raw_result, agent="parser", user=raw_text[:500])
    elapsed = int((time.time() - t0) * 1000)
    logger.info(f"CV parsed in {elapsed}ms: {result.get('full_name', 'Unknown')}")

    if not result:
        return {"full_name": "Unknown", "email": None, "phone": None}

    # Normalize phone: sometimes LLM returns a list
    phone = result.get("phone")
    if isinstance(phone, list):
        result["phone"] = phone[0] if phone else None
    if isinstance(result.get("years_experience"), str):
        try:
            result["years_experience"] = float(result["years_experience"].split()[0])
        except Exception:
            result["years_experience"] = 0.0

    return result


def match_cv_to_job(parsed_cv: dict, job: dict) -> dict:
    """
    Score candidate against specific job.
    CRITICAL: result is strictly for THIS candidate + THIS job.
    Never cache or reuse across different (candidate, job) pairs.
    """
    import json as _json
    t0 = time.time()

    user_prompt = (
        f"JOB REQUIREMENTS:\n"
        f"Title: {job.get('title', '')}\n"
        f"Description: {job.get('description', '')[:3000]}\n"
        f"Required Skills: {job.get('required_skills', [])}\n"
        f"Nice to Have: {job.get('nice_to_have', [])}\n"
        f"Min Experience: {job.get('min_experience', 0)} years\n"
        f"Education: {job.get('education_req', 'Any')}\n"
        f"Location: {job.get('location_req', 'Any')}\n\n"
        f"CANDIDATE PROFILE:\n"
        f"{_json.dumps(parsed_cv, ensure_ascii=False)[:4000]}"
    )

    result = llm_complete_json(
        system=MATCH_SYSTEM,
        user=user_prompt,
        max_tokens=1500,
        agent="matcher",
    )
    elapsed = int((time.time() - t0) * 1000)
    score = result.get("overall_score", 0)
    logger.info(f"Job matched in {elapsed}ms: score={score}")

    if not result:
        return _empty_match()

    # Ensure score is float 0-100
    for key in ("overall_score", "skill_match", "experience_match", "education_match",
                "seniority_match", "location_match", "keyword_match", "ats_score", "ai_confidence"):
        v = result.get(key, 0)
        try:
            result[key] = float(v)
        except Exception:
            result[key] = 0.0

    return result


def _empty_match() -> dict:
    return {
        "overall_score": 0, "skill_match": 0, "experience_match": 0,
        "education_match": 0, "seniority_match": 0, "location_match": 0,
        "keyword_match": 0, "ats_score": 0, "ai_confidence": 0,
        "recommendation": "Consider",
        "recommendation_reason": "Analysis unavailable",
        "ai_summary": "", "strengths": [], "weaknesses": [],
        "missing_skills": [], "missing_certs": [], "matched_skills": [],
        "matched_requirements": [], "missing_requirements": [],
        "skill_gap_analysis": "", "ats_issues": [], "ats_suggestions": [],
    }
